import io
import difflib
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS  
from transformers import pipeline, AutoModelForSeq2SeqLM, AutoTokenizer
from docx import Document
from docx.shared import RGBColor
import PyPDF2

app = Flask(__name__)
CORS(app)  

# Load AI model
model_name = "facebook/bart-large-cnn"
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
comparison_model = pipeline("summarization", model=model, tokenizer=tokenizer)

def extract_text_from_doc(file, file_type):
    """Extract text from a DOCX or PDF file."""
    if file_type == "docx":
        doc = Document(io.BytesIO(file.read()))
        return " ".join([para.text for para in doc.paragraphs])
    elif file_type == "pdf":
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
        text = "".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
        return text.strip()
    return ""

def highlight_differences(original, edited):
    """Highlight changes in the documents."""
    original_words = original.split()
    edited_words = edited.split()
    diff = list(difflib.ndiff(original_words, edited_words))
    
    original_highlighted = []
    edited_highlighted = []
    
    for word in diff:
        if word.startswith('- '):  # Deleted word
            original_highlighted.append(f"<del>{word[2:].strip()}</del>")
        elif word.startswith('+ '):  # Added word
            edited_highlighted.append(f"<ins>{word[2:].strip()}</ins>")
        else:
            original_highlighted.append(word[2:].strip())
            edited_highlighted.append(word[2:].strip())
    
    return " ".join(original_highlighted), " ".join(edited_highlighted)

def summarize_text(text):
    """Generate a summary of the given text."""
    if len(text.split()) < 50:
        return text
    summary = comparison_model(text, max_length=100, min_length=30, truncation=True)
    return summary[0]['summary_text']

def create_highlighted_docx(original_highlighted, edited_highlighted):
    """Generate a properly formatted DOCX file with text styling instead of raw HTML."""
    doc = Document()
    doc.add_heading("Document Comparison Result", level=1)

    # Process Original Document with Deleted Text in Red
    doc.add_heading("Original Document (Deletions in Red)", level=2)
    para = doc.add_paragraph()
    for word in original_highlighted.split():
        if word.startswith("<del>") and word.endswith("</del>"):
            run = para.add_run(word[5:-6])  # Remove <del> tags
            run.font.color.rgb = RGBColor(255, 0, 0)  # Set color to red
            run.bold = True
        else:
            para.add_run(word + " ")  # Normal text

    # Process Edited Document with Inserted Text in Green
    doc.add_heading("Edited Document (Additions in Green)", level=2)
    para = doc.add_paragraph()
    for word in edited_highlighted.split():
        if word.startswith("<ins>") and word.endswith("</ins>"):
            run = para.add_run(word[5:-6])  # Remove <ins> tags
            run.font.color.rgb = RGBColor(0, 128, 0)  # Set color to green
            run.bold = True
        else:
            para.add_run(word + " ")  # Normal text

    file_path = "highlighted_comparison.docx"
    doc.save(file_path)
    return file_path

@app.route('/compare', methods=['POST'])
def compare_documents():
    file1 = request.files.get('file1')
    file2 = request.files.get('file2')

    if not file1 or not file2:
        return jsonify({"error": "Both files are required."}), 400

    file1_type = file1.filename.split('.')[-1].lower()
    file2_type = file2.filename.split('.')[-1].lower()
    
    if file1_type not in ['docx', 'pdf'] or file2_type not in ['docx', 'pdf']:
        return jsonify({"error": "Only DOCX and PDF files are supported."}), 400

    file1_content = extract_text_from_doc(file1, file1_type)
    file2_content = extract_text_from_doc(file2, file2_type)

    original_highlighted, edited_highlighted = highlight_differences(file1_content, file2_content)
    original_summary = summarize_text(file1_content)
    edited_summary = summarize_text(file2_content)

    highlighted_doc_path = create_highlighted_docx(original_highlighted, edited_highlighted)

    return jsonify({
        "original_highlighted": original_highlighted,
        "edited_highlighted": edited_highlighted,
        "original_summary": original_summary,
        "edited_summary": edited_summary,
        "download_url": "/download"  # URL to download the highlighted DOCX
    })

@app.route('/download', methods=['GET'])
def download_highlighted_docx():
    """Provide the highlighted DOCX file for download."""
    file_path = "highlighted_comparison.docx"
    return send_file(file_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
